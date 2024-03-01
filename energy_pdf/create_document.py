import json
import pathlib
import uuid
from copy import deepcopy
import subprocess
import docx
from docxtpl import DocxTemplate
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# in those three need . before file like .charts
from .charts import create_bar_chart, create_donut_chart, create_history_chart
from .utils.add_table_data import *
from .utils.placeholders import get_placeholders_from_invoice

import os
import shutil


class EnergyPdfGenerator:
    language: str = None
    invoice_data = None
    doc = None
    tables_map = None
    client_data: dict | None = None
    yaml_file_path = str(pathlib.Path(__file__).parent / 'utils/placeholders.yaml')
    template_filepath_it = str(pathlib.Path(__file__).parent / "templates/energy_template_it.docx")
    template_filepath_de = str(pathlib.Path(__file__).parent / "templates/energy_template_de.docx")
    docx_temp_folder = f"/tmp/{uuid.uuid4()}"
    docx_temp_filepath = f"{docx_temp_folder}/document.docx"
    pdf_temp_filepath = f"{docx_temp_folder}/document.pdf"
    dynamic_tables = [
        'SPESA_PER_LA_MATERIA',
        'SPESA_PER_IL_TRASPORTO',
        'SPESA_PER_ONERI_DI_SISTEMA',
        'IMPOSTE'
    ]

    def __init__(self, invoice_data: dict, language: str, client_data):
        self.invoice_data = invoice_data
        self.language = language
        self.client_data = client_data
        pathlib.Path(self.docx_temp_folder).mkdir(parents=True, exist_ok=True)

    def set_repeat_table_header(self, row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)

    def add_new_row_and_fill_data(self, table, source_row, data: list, header_row=None):
        if header_row:
            self.set_repeat_table_header(header_row)
        if not source_row:
            return
        for _ in range(len(data)):
            new_row_element = deepcopy(source_row._element)
            new_row = table.add_row()
            new_row._element.getparent().replace(new_row._element, new_row_element)

        i = 0
        for index in range(len(table.rows) - len(data) - 1, len(table.rows) - 1):
            row = table.rows[index]
            for index_of_cell, cell in enumerate(row.cells):
                if index_of_cell == len(data[i]):
                    break

                cell.paragraphs[0].text = str(data[i][index_of_cell])

            i += 1

        self.delete_identification_row(table)

    def add_data_to_dynamic_tables(self):

        for table in self.dynamic_tables:
            table_docx = self.tables_map[table]['table']
            row_docx = self.tables_map[table]['model_row']

            if table in ['SPESA_PER_LA_MATERIA', 'SPESA_PER_IL_TRASPORTO', 'SPESA_PER_ONERI_DI_SISTEMA', 'IMPOSTE']:
                json_detail = str(self.tables_map[table].get('json'))
                header_row = self.tables_map[table].get('header_row')
                self.add_data_to_details_tables(table_docx, row_docx, json_detail, header_row)

    def add_data_to_details_tables(self, table_docx, row_docx, json_detail: str, header_row=None):
        detail_data = self.invoice_data['invoice_json']['detail'][json_detail]

        insert_data = []
        for item in detail_data:
            lettura = ''
            if item.get('estimated'):
                if self.language == 'it':
                    lettura = 'stimate'
                else:
                    lettura = 'geschätzt'
            else:
                if self.language == 'it':
                    lettura = 'effetive'
                else:
                    lettura = 'tatsächlich'

            insert_data.append([
                item['label'],
                f'{self.format_date(item["period_start"])} - {self.format_date(item["period_end"])}',
                str(item['unit_type']).replace('EUR', '€').replace('_per_', '/').replace('_and_', '&'),
                str(item['unit_price']).replace('.', ','),
                str(item['unit']),
                str(item['total_price']),
                f'{lettura}',
                f'{str(int(item["vat_percentage"] * 100))}%'
            ])

        self.add_new_row_and_fill_data(table_docx, row_docx, insert_data, header_row)

    def add_data_to_specific_tables(self):
        dettaglio_data = get_data_for_dettaglio_fiscale(self.invoice_data)
        self.add_new_row_and_fill_data(self.tables_map['DETTAGLIO_FISCALE']['table'], self.tables_map['DETTAGLIO_FISCALE']['model_row'], dettaglio_data)
        dettaglio_iva = get_data_for_dettaglio_iva(self.invoice_data)
        self.add_new_row_and_fill_data(self.tables_map['DETTAGLIO_IVA']['table'], self.tables_map['DETTAGLIO_IVA']['model_row'], dettaglio_iva)
        letture_data = get_data_for_letture(self.invoice_data)
        self.add_new_row_and_fill_data(self.tables_map['LETTURE']['table'], self.tables_map['LETTURE']['model_row'], letture_data)
        iva_last_page = get_data_for_iva_last_page(self.invoice_data)
        self.add_new_row_and_fill_data(self.tables_map['IVA']['table'], self.tables_map['IVA']['model_row'], iva_last_page)
        sintesi_data = get_data_for_sintesi(self.invoice_data)
        self.add_new_row_and_fill_data(self.tables_map['SINTESI_IMPORTI_E_FATTURATI']['table'], self.tables_map['SINTESI_IMPORTI_E_FATTURATI']['model_row'], sintesi_data)
        livelo_massimo_data = get_data_for_livelo_massimo(self.invoice_data)
        self.add_new_row_and_fill_data(self.tables_map['LIVELO_MASSIMO_DI_POTENZA_PRELEVATA']['table'], self.tables_map['LIVELO_MASSIMO_DI_POTENZA_PRELEVATA']['model_row'],
                                       livelo_massimo_data)

    def format_date(self, date):
        date_object = datetime.strptime(date, "%Y-%m-%d")
        formatted_string = date_object.strftime("%d/%m/%Y")

        return formatted_string

    def delete_identification_row(self, table):
        last_row = table.rows[-1]

        tr = last_row._element

        tbl = tr.getparent()

        tbl.remove(tr)

    def populate_placeholders(self):
        placeholders = get_placeholders_from_invoice(self.invoice_data, self.yaml_file_path, self.language)
        placeholders.update(self.client_data)
        doc = DocxTemplate(self.docx_temp_filepath)
        doc.render(placeholders)
        doc.save(self.docx_temp_filepath)

    def replace_image(self, new_image_path, old_image):
        try:
            for rel_key, rel in self.doc.part.rels.items():
                # morao sam da hardkoriram ime slike jer joj docx menja ime, bar kada sacuvam preko libreoffice-a
                if rel.reltype and old_image in rel.target_ref:
                    # Replace the image
                    with open(new_image_path, 'rb') as new_img:
                        rel.target_part._blob = new_img.read()
                    break
        except Exception as e:
            raise

    def convert_to_pdf(self) -> None:
        unique_profile_dir = os.path.join("/tmp", f"libreoffice_profile_{uuid.uuid4()}")
        profile_uri = f"file://{unique_profile_dir}"  # Convert the path to a URI

        cmd = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            # "soffice",
            '--headless',
            '--convert-to', 'pdf',
            self.docx_temp_filepath,
            '--outdir', self.docx_temp_folder,
            '-env:UserInstallation=' + profile_uri
        ]

        try:
            subprocess.run(cmd, check=True)
        except subprocess.CalledProcessError as e:
            print("An error occurred while converting the file:", e)
            raise
        except Exception as e:
            print("An unexpected error occurred:", e)
            raise
        finally:
            # Clean up: Remove the user profile directory after conversion
            if os.path.exists(unique_profile_dir):
                shutil.rmtree(unique_profile_dir)

    def generate_pdf(self) -> tuple:
        if self.language == 'it':
            self.doc = docx.Document(self.template_filepath_it)
        elif self.language == 'de':
            self.doc = docx.Document(self.template_filepath_de)
        doc = self.doc
        self.tables_map = {
            'SINTESI_IMPORTI_E_FATTURATI': {
                'table': doc.tables[0],
                'model_row': doc.tables[0].rows[2]
            },

            'DETTAGLIO_FISCALE': {
                'table': doc.tables[4],
                'model_row': doc.tables[4].rows[1]
            },
            'DETTAGLIO_IVA': {
                'table': doc.tables[5],
                'model_row': doc.tables[5].rows[1]
            },
            'TOTALE_BOLLETA_1': {
                'table': doc.tables[6],
                'model_row': doc.tables[6].rows[2]
            },
            'TOTALE_DA_PAGARE_1': {
                'table': doc.tables[7],
                'model_row': doc.tables[7].rows[1]
            },
            'LETTURE': {
                'table': doc.tables[8],
                'model_row': doc.tables[8].rows[3]
            },
            'LIVELO_MASSIMO_DI_POTENZA_PRELEVATA': {
                'table': doc.tables[9],
                'model_row': doc.tables[9].rows[3]
            },
            'SPESA_PER_LA_MATERIA': {
                'table': doc.tables[15],
                'model_row': doc.tables[15].rows[3],
                'header_row': doc.tables[15].rows[2],
                'json': 'energy'
            },
            'SPESA_PER_IL_TRASPORTO': {
                'table': doc.tables[17],
                'model_row': doc.tables[17].rows[1],
                'header_row': doc.tables[17].rows[0],
                'json': 'transport'
            },
            'SPESA_PER_ONERI_DI_SISTEMA': {
                'table': doc.tables[19],
                'model_row': doc.tables[19].rows[1],
                'header_row': doc.tables[19].rows[0],
                'json': 'system'
            },
            'IMPOSTE': {
                'table': doc.tables[21],
                'model_row': doc.tables[21].rows[1],
                'header_row': doc.tables[21].rows[0],
                'json': 'tax'
            },
            'IVA': {
                'table': doc.tables[23],
                'model_row': doc.tables[23].rows[1],
                'header_row': doc.tables[23].rows[0]
            }
        }
        self.add_data_to_dynamic_tables()
        self.add_data_to_specific_tables()
        bar_chart_filepath = create_bar_chart(self.invoice_data)
        donut_chart_filepath = create_donut_chart(self.invoice_data)
        hist_chart_filepath = create_history_chart(self.invoice_data)
        self.replace_image(bar_chart_filepath, "media/image7.png")
        self.replace_image(donut_chart_filepath, "media/image8.png")
        self.replace_image(hist_chart_filepath, "media/image10.png")
        self.doc.save(self.docx_temp_filepath)
        self.populate_placeholders()

        self.convert_to_pdf()
        if os.path.exists(bar_chart_filepath):
            os.remove(bar_chart_filepath)
        if os.path.exists(donut_chart_filepath):
            os.remove(donut_chart_filepath)
        if os.path.exists(hist_chart_filepath):
            os.remove(hist_chart_filepath)
        return self.docx_temp_folder, self.pdf_temp_filepath
